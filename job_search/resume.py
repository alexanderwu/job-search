"""
Convert markdown resume or job listings to docx and pdfs

Usage:
>>> import resume as res
>>> P_resume = "Alexander_Wu_Resume.md"
>>> convert_resume(P_resume)
"""
from functools import cache
from pathlib import Path

import docx
import docx2pdf
import pandas as pd
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from mdit_py_plugins.front_matter import front_matter_plugin


P_ALEX_RESUME_MD = Path('data/Alexander_Wu_Resume.md')
P_RAW = Path('data/raw')


def convert_markdown(path_md: Path | str, keep_docx=True, resume=False, pagebreak=False, verbose=True):
    P_md = Path(path_md)
    mdf = analyze(P_md, explode=True)
    mdf['_style'] = 'Normal'
    mdf.loc[lambda x: (x['_tag'] == 'h2') & (x['_i'] == 0), '_style'] = 'Heading 2'
    mdf.loc[lambda x: x['_tag'].isin(['ul', 'li']) & (x['_i'] == 0), '_style'] = 'List Bullet'
    if pagebreak:
        _pagebreak_list = ((mdf['_tag'] == "ul") & (mdf['_markdown'].str.len() == 0)).to_list()
        _first_index = _pagebreak_list.index(True)
        _second_index = _pagebreak_list.index(True, _first_index + 1)
        mdf.loc[_second_index, '_style'] = '_pagebreak'
    if resume:
        mdf.loc[0, '_style'] = 'title'
        mdf.loc[1, '_style'] = 'subtitle'
        mdf.loc[3, '_style'] = 'Heading 2'
    _bullets = (' ' + mdf['_condensed']).str.extractall(r'(.+?\*\*.+?\*\*)').groupby(level=0)[0].apply(list)
    _bullets = _bullets.reindex_like(mdf).apply(lambda x: x if isinstance(x, list) else [])
    _last_item = ('**' + mdf['_condensed']).str.split('**', regex=False).str[-1]
    mdf['_runs'] = _bullets + _last_item.fillna('').apply(lambda x: [x])

    document = docx.Document()
    _init_document(document, header=resume)

    for _, row in mdf.dropna(subset='_condensed').iterrows():
        # _string = ''.join(row['_runs']).strip()
        # print(f" {row['_line']:2}: {row['_style']:12} - {row['_runs']} =={_string}==")
        _add_paragraph(document, row['_runs'], row['_style'])

    P_docx = P_md.parent / f"{P_md.stem}"
    # P_docx = Path('data/_') / f"{P_md.stem}"
    if verbose:
        print(f'Saving to {P_docx}')
    document.save(f'{P_docx}.docx')
    convert_pdf(document, P_docx, keep_docx=keep_docx)


def convert_resume(path_md: Path | str = P_ALEX_RESUME_MD, keep_docx=False, pagebreak=False, verbose=True):
    convert_markdown(path_md, keep_docx, resume=True, pagebreak=pagebreak, verbose=verbose)


@cache
def analyze(path_md: Path | str = P_ALEX_RESUME_MD, explode=True) -> pd.DataFrame:
    """
    Returns:
        mdf (pd.DataFrame): markdown dataframe representation
    """
    _md = MarkdownIt('commonmark', {'breaks': True, 'html': True}).use(front_matter_plugin)
    with open(path_md, encoding='utf-8') as f:
        # _md_text = f.read()
        _md_lines = f.readlines()
        _md_text = ''.join([line for line in _md_lines if not line.startswith("<!--")])
        tokens = _md.parse(_md_text)

    tdf = pd.DataFrame(tokens)
    tdf['_begin'] = tdf['map'].ffill().str[0]

    mdf = tdf.groupby('_begin').agg(
        _len=('_begin', len),  # pl.len(),
        _tag=('tag', 'last'),  # pl.col('tag').last(),
        markup=('markup', 'first'),  # pl.col('markup').first(),
        content=('content', ''.join),  # pl.col('content').implode().list.join(''),
    ).reset_index()
    # mdf['_tag'] = mdf['_tag'] + (mdf['_tag'] == mdf['_tag'].shift(-1)).map({True: '_', False: ''})
    mdf['_ul_next'] = mdf['_tag'].shift(-1).isin(['ul', 'li'])

    _md_suffix_newline = mdf['_tag'].isin(['li']).map({False: '\n', True: ''})
    _markdown_line = (mdf['markup'] + ' ' + mdf['content']).str.lstrip() + _md_suffix_newline
    mdf['_markdown'] = _markdown_line.str.split('\n')
    _suffix_newline = (mdf['_tag'].isin(['li', 'h2']) | mdf['_ul_next']).map({False: '\n', True: ''})
    _condensed_line = mdf['content'] + _suffix_newline
    mdf['_condensed'] = _condensed_line.str.split('\n')

    if explode:
        md_df = mdf.explode('_markdown').drop(columns='content').reset_index(drop=True)
        md_df['_line'] = 1 + md_df['_begin'] + md_df.groupby('_begin').cumcount()
        md_df['_i'] = md_df.groupby('_begin').cumcount()
        md_df['_condensed'] = md_df.apply(lambda x: x['_condensed'][x['_i']] if x['_i'] < len(x['_condensed']) else None, axis=1)
        return md_df
    return mdf


def convert_pdf(document: Document, name: Path | str, keep_docx=False):
    """Converts a DOCX file to PDF using LibreOffice command line."""
    import subprocess

    # docx2pdf.convert(f'{name}.docx', f'{name}.pdf')
    command = ["swriter", "--headless", "--convert-to", "pdf", "--outdir", P_RAW, name]
    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15)
        print(f"Successfully converted {name} to PDF.")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e.stderr.decode()}")
    except FileNotFoundError:
        print("LibreOffice not found. Make sure it is installed and in your PATH.")
    except subprocess.TimeoutExpired:
        print("Conversion timed out.")

    if not keep_docx:
        Path(f'{name}.docx').unlink()


def _init_document(document, header=True):
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    # section.left_margin = Inches(0.7)
    # section.right_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    if header:
        section.different_first_page_header_footer = True
        pp = section.header.paragraphs[0]
        pp.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)
        pp.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.CENTER)
        pp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        rr = pp.add_run("\tAlexander Wu\talexander.wu7@gmail.com\tPage 2\n")
        rr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # _add_page_number(rr)

    normal_style = document.styles['Normal']
    normal_style.font.name = 'Calibri'
    # normal_style.font.size = Pt(12)
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # normal_style.paragraph_format.line_spacing = 1.05
    normal_style.paragraph_format.space_before = 0
    normal_style.paragraph_format.space_after = 0
    normal_style.paragraph_format.tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)

    title_style = document.styles['Title']
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.font.size = Pt(18)
    subtitle_style = document.styles['Subtitle']
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.font.size = Pt(11)

    h2_style = document.styles['Heading 2']
    # h2_style.font.color.rgb = RGBColor(0x3d, 0x59, 0x7f)
    h2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2_style.font.size = Pt(14)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # h2_style.paragraph_format.space_before = 0

    ul_style = document.styles['List Bullet']
    ul_style.paragraph_format.left_indent = Inches(0.5)
    ul_style.paragraph_format.space_before = Pt(2)


def _add_paragraph(document, runs, style='Normal'):
    _string = ''.join(runs).strip()
    if _string.startswith("<!--"):  # <!-- comment -->
        return

    if style == "_pagebreak":
        document.add_page_break()
    elif style == "title":
        _add_title(document, _string.strip('*'))
    elif style == "subtitle":
        _add_subtitle(document, _string)
    else:
        pp = document.add_paragraph(style=style)
        for run in runs[:-1]:
            _normal, _bold, *_ = run.split('**')
            if len(_normal := _normal.lstrip()) > 0:
                pp.add_run(_normal)
            if len(_bold) > 0:
                pp.add_run(_bold).bold = True
        if len(_normal := runs[-1]) > 0:
            pp.add_run(_normal)


def _add_title(document, title):
    pp = document.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(title)
    rr.font.size = Pt(18)
    rr.bold = True

def _add_subtitle(document, subtitle):
    pp = document.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(subtitle)
    rr.font.size = Pt(11)
    rr.bold = False

def _add_page_number(run, separate=True):
    """

    <w:fldChar w:fldCharType="begin"/>
    <w:instrText xml:space="preserve">PAGE</w:instrText>
    <w:fldChar w:fldCharType="separate"/>
    <w:fldChar w:fldCharType="end"/>
    <w:br/>
    """
    _page_xml = _create_xml('w:instrText', 'xml:space', 'preserve')
    _page_xml.text = "PAGE"
    run._r.append(_create_xml('w:fldChar', 'w:fldCharType', 'begin'))
    run._r.append(_page_xml)
    if separate:
        run._r.append(_create_xml('w:fldChar', 'w:fldCharType', 'separate'))
    run._r.append(_create_xml('w:fldChar', 'w:fldCharType', 'end'))
    run._r.append(_create_xml('w:br'))
    # run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def _create_xml(name, attr=None, val=None):
    element = OxmlElement(name)
    if attr is not None and val is not None:
        element.set(ns.qn(attr), val)
    return element


if __name__ == "__main__":
    # from job_search.config import P_RAW
    # P_resume = P_ALEX_RESUME_MD  # Alexander_Wu_Resume.md
    # P_resume_pdf = P_resume.parent / f"{P_resume.stem}.pdf"
    # print(f"Converting {P_resume} to {P_resume_pdf}...")
    # convert_resume(P_resume, keep_docx=False)

    # P_beone_resume = Path('data/raw') / 'Alexander_Wu_Resume - BeOne Medicines.md'
    # P_resume = P_RAW / 'AW_Roche_Resume.md'
    # P_resume = Path('data/raw') / 'AW_Healthcare_Resume.md'
    # P_resume = Path('data/raw') / 'Alex_Wu_Resume - RWD Programmer.md'
    # P_resume = Path('data/raw') / 'Alex_Wu_Verily_Resume.md'
    P_resume = Path('data/raw') / 'AW_Resume.md'
    # P_resume = Path('data/interim') / 'Alex_Wu_Resume - Regeneron.md'
    # P_resume = Path('data/interim') / 'Alex_Wu_Capegemini_Resume.md'
    # P_resume = Path('data/interim') / 'AW_Genentech_MDAE_Resume.md'
    # P_resume = Path('data/interim') / 'AW_Homeward_Resume.md'
    # P_resume = Path('data/interim') / 'AW_Oscar_Resume.md'
    # P_resume = Path('data/interim') / 'AW_Verily_Resume.md'
    # convert_resume(P_resume, keep_docx=False, pagebreak=True)
    # convert_resume(P_resume, keep_docx=True, pagebreak=True)
    convert_resume(P_resume, keep_docx=True, pagebreak=True)
    #swriter --headless --convert-to pdf "data/interim/AW_Genentech_MDAE_Resume.docx" --outdir data/processed/
    #swriter --headless --convert-to pdf "data/interim/AW_Homeward_Resume.docx" --outdir data/processed/
    #swriter --headless --convert-to pdf "data/interim/AW_Oscar_Resume.docx" --outdir data/processed/
    #swriter --headless --convert-to pdf "data/interim/AW_Verily_Resume.docx" --outdir data/processed/
    #swriter --headless --convert-to pdf "data/raw/AW_Resume.docx" --outdir data/processed/
    #code data/processed/AW_Verily_Resume.pdf
